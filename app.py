import os
import json
import requests
import re
import difflib
import io
import docx
import openpyxl
import pdfplumber
import sqlite3
import time
import uuid
import urllib.parse
import tempfile
import logging
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from dotenv import load_dotenv
from fastapi.staticfiles import StaticFiles
from urllib.parse import quote
from datetime import datetime, timezone, timedelta
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel
from typing import List
from openai import OpenAI

# 设置中国时区（UTC+8）
CHINA_TZ = timezone(timedelta(hours=8))

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)

app = FastAPI(title="ISO 体系智能审核系统 API")

# 加载环境变量
load_dotenv()

# --- 新增：开启静态文件服务，让 XDOC 能读取到原文件 ---
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)
app.mount("/uploads", StaticFiles(directory=UPLOAD_DIR), name="uploads")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- 配置区（从环境变量读取） ---
TEXTIN_APP_ID = os.getenv("TEXTIN_APP_ID")
TEXTIN_SECRET_CODE = os.getenv("TEXTIN_SECRET_CODE")
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")
ZHIPU_API_KEY = os.getenv("ZHIPU_API_KEY")

if not all([TEXTIN_APP_ID, TEXTIN_SECRET_CODE, DEEPSEEK_API_KEY]):
    print("警告: 请在 .env 文件中配置所有 API 密钥")

# DeepSeek 客户端（用于审核）
deepseek_client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com/v1")

# 智谱 AI 客户端（用于提示词优化）
zhipu_client = OpenAI(api_key=ZHIPU_API_KEY, base_url="https://ark.cn-beijing.volces.com/api/coding/v3")

DATA_DIR = os.getenv("DATA_DIR", "data")
DB_FILE = os.getenv("DB_FILE", "audit_batches.db")

# 使用绝对路径，确保在任何工作目录下都能找到
if not os.path.isabs(DATA_DIR):
    DATA_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), DATA_DIR)
if not os.path.isabs(DB_FILE):
    DB_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), DB_FILE)

BACKEND_HOST = os.getenv("BACKEND_HOST", "0.0.0.0")
BACKEND_PORT = int(os.getenv("BACKEND_PORT", 8000))

with open("element_routing_dictionary_final.json", "r", encoding="utf-8") as f:
    routing_dict = json.load(f)

# --- 数据库初始化 ---
def init_db():
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    # 历史报告表
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS batch_reports (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            batch_name TEXT,
            total_count INTEGER,
            pass_count INTEGER,
            pass_rate REAL,
            details_json TEXT,
            created_at DATETIME DEFAULT (datetime('now', '+8 hours'))
        )
    ''')

    # 新增：系统提示词表
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS system_prompts (
            prompt_key TEXT PRIMARY KEY,
            prompt_name TEXT,
            content TEXT,
            updated_at DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')

    # 初始化默认提示词（如果表为空）
    cursor.execute('SELECT count(*) FROM system_prompts')
    if cursor.fetchone()[0] == 0:
        default_standard = """你是一个拥有20年经验的ISO体系高级内审员，你的审核风格是：逻辑极其严密、注重实质合规、且具备极强的“单文件深度分析”能力。
【审核标准文件原文】（判罚依据）：\n{standard_text}
【待审核证据文件内容】：\n{evidence_text}
请根据【审核标准文件原文】，对【待审核证据文件内容】进行深度的内容逻辑与合规性审核。
（请保持严格的JSON输出：包含"结论"、"逻辑缺陷描述"、"引用依据"、"修改意见"）"""

        default_daily = """你是一个资深的EHS安全专家和卓越的数据分析师。
【8.1要素管理标准原文】：\n{standard_text}
【待分析的日常检查汇总表内容】：\n{evidence_text}
系统已统计：总数{total}，已完成{completed}，整改率{rate}%，前三类别{top3_str}。
请结合上述数据和标准原文，固定输出"符合"结论，并给出专业的分析和整改意见。"""

        cursor.execute("INSERT INTO system_prompts (prompt_key, prompt_name, content) VALUES (?, ?, ?)",
                       ('standard_audit', '标准体系交叉审核提示词', default_standard))
        cursor.execute("INSERT INTO system_prompts (prompt_key, prompt_name, content) VALUES (?, ?, ?)",
                       ('daily_inspection', '日常检查台账分析提示词', default_daily))

    conn.commit()
    conn.close()

init_db()

# --- 提示词数据库操作 ---
def get_prompt_from_db(prompt_key: str) -> str:
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute("SELECT content FROM system_prompts WHERE prompt_key = ?", (prompt_key,))
    row = cursor.fetchone()
    conn.close()
    return row[0] if row else ""

# 检查数据库和数据目录
logger.info(f"数据目录: {DATA_DIR}")
logger.info(f"数据库文件: {DB_FILE}")
logger.info(f"上传目录: {UPLOAD_DIR}")

# --- 【新增核心模块】：Word 文档顺序遍历器 ---
def iter_block_items(parent):
    """
    深度解析 Word 底层 XML 树，严格按照从上到下的物理顺序生成段落和表格。
    """
    if isinstance(parent, DocxDocument):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("不支持的节点类型")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def parse_docx_in_order(file_stream) -> str:
    """顺序解析 Docx 为 Markdown"""
    doc = docx.Document(file_stream)
    content = []
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.replace('\n', ' ').strip()
            if text:
                content.append(text)
        elif isinstance(block, Table):
            content.append("\n")
            for i, row in enumerate(block.rows):
                row_data = [cell.text.replace('\n', ' ').replace('\r', '').strip() for cell in row.cells]
                content.append("| " + " | ".join(row_data) + " |")
                if i == 0:
                    content.append("|" + "|".join(["---"] * len(row.cells)) + "|")
            content.append("\n")
    return "\n".join(content)

def parse_pdf_in_order(pdf_stream) -> tuple:
    """顺序解析 PDF，并返回(文本内容, 是否包含大量图片需转交OCR)"""
    content = []
    requires_ocr = False

    with pdfplumber.open(pdf_stream) as pdf:
        for page in pdf.pages:
            # 1. 检查图片占比，决定是否触发 TextIn 云端兜底
            page_area = float(page.width * page.height)
            for img in page.images:
                img_area = float(img.get('width', 0) * img.get('height', 0))
                if page_area > 0 and (img_area / page_area) > 0.05:
                    requires_ocr = True
                    break

            # 2. 提取页面上所有的表格，并记录它们的包围盒(bbox)
            tables = page.find_tables()
            table_bboxes = [t.bbox for t in tables] # [x0, top, x1, bottom]

            # 3. 提取所有的文字单词，附带它们的坐标
            words = page.extract_words(keep_blank_chars=True)

            # 4. 过滤掉那些"在表格内部"的文字（因为表格我们要转成 Markdown，不能重复提取）
            non_table_words = []
            for w in words:
                in_table = False
                for bbox in table_bboxes:
                    if bbox[0] <= w['x0'] and bbox[1] <= w['top'] and bbox[2] >= w['x1'] and bbox[3] >= w['bottom']:
                        in_table = True
                        break
                if not in_table:
                    non_table_words.append(w)

            # 5. 将纯文本和 Markdown 表格封装成具有 `top` (Y坐标) 的对象，扔进一个大列表里
            page_elements = []

            # 合并同一行的文字
            current_line_text = ""
            current_line_top = -1
            line_tolerance = 5 # Y 坐标误差 5 内视为同一行

            for w in non_table_words:
                if current_line_top == -1 or abs(w['top'] - current_line_top) < line_tolerance:
                    current_line_text += w['text']
                    if current_line_top == -1: current_line_top = w['top']
                else:
                    page_elements.append({'type': 'text', 'top': current_line_top, 'content': current_line_text.strip()})
                    current_line_text = w['text']
                    current_line_top = w['top']
            if current_line_text:
                page_elements.append({'type': 'text', 'top': current_line_top, 'content': current_line_text.strip()})

            # 将表格转化为 Markdown 并加入列表
            for t_obj in tables:
                table_content = ["\n"]
                table_data = t_obj.extract()
                for i, row in enumerate(table_data):
                    clean_row = [str(c).replace('\n', ' ').replace('\r', '').strip() if c else "" for c in row]
                    table_content.append("| " + " | ".join(clean_row) + " |")
                    if i == 0:
                        table_content.append("|" + "|".join(["---"] * len(clean_row)) + "|")
                table_content.append("\n")

                page_elements.append({
                    'type': 'table',
                    'top': t_obj.bbox[1],
                    'content': "\n".join(table_content)
                })

            # 6. 【最关键的一步】：按照 Y 坐标 (top) 从上到下对文字和表格进行排序！
            page_elements.sort(key=lambda x: x['top'])

            # 7. 拼装成本页的最终内容
            for el in page_elements:
                if el['content']:
                    content.append(el['content'])

    return "\n".join(content).strip(), requires_ocr


def read_local_file_content(file_path: str) -> str:
    """通用本地文件读取函数（升级版：自上而下顺序解析）"""
    if not os.path.exists(file_path):
        return ""

    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext == ".docx":
            with open(file_path, "rb") as f:
                return parse_docx_in_order(f)
        elif ext == ".pdf":
            with open(file_path, "rb") as f:
                text, _ = parse_pdf_in_order(f)
                return text
    except Exception as e:
        print(f"解析标准文件 {file_path} 失败: {e}")
    return ""

def get_standard_document_content(major_element: str, matched_rule: str) -> str:
    """
    根据匹配到的规则名（无论是子要素还是兜底的主要素文件）去读取标准原文
    """
    if matched_rule:
        for ext in [".docx", ".pdf", ".doc"]:
            path = os.path.join(DATA_DIR, f"{matched_rule}{ext}")
            content = read_local_file_content(path)
            if content:
                return f"【审核标准原文：{matched_rule}】\n{content}"

    # 极限兜底：如果连文件都没建好，直接给提示
    return "未在 data 文件夹下找到对应的标准原文文件，请仅根据通用ISO常识审核。"

def extract_text_via_textin(file_bytes: bytes) -> str:
    """调用 TextIn xParse 文档解析接口，直接输出原生 Markdown"""
    # 【改动1：URL 换成了专门转 Markdown 的大模型接口】
    url = "https://api.textin.com/ai/service/v1/pdf_to_markdown"

    headers = {
        "x-ti-app-id": TEXTIN_APP_ID,
        "x-ti-secret-code": TEXTIN_SECRET_CODE,
        "Content-Type": "application/octet-stream" # 官方文档要求的格式
    }

    # 【改动2：增加 URL 控制参数】
    # table_flavor="md" 告诉引擎：一定要把图片里的表格转成纯正的 Markdown 分割线！
    params = {
        "parse_mode": "auto",
        "table_flavor": "md",
        "apply_document_tree": 1
    }

    try:
        # 发送请求，注意这里多传了一个 params
        response = requests.post(url, headers=headers, params=params, data=file_bytes)
        result = response.json()

        # 【改动3：直接从 result 字典里提取 markdown 字段】
        if result.get("code") == 200:
            return result.get("result", {}).get("markdown", "")
        else:
            print(f"TextIn 接口报错: {result.get('message')}")
    except Exception as e:
        print(f"云端解析异常: {e}")
        pass

    return "解析失败"

def extract_upload_file_text(file_bytes: bytes, filename: str) -> str:
    """解析用户上传的待审证据文件：自上而下顺序解析 + 云端兜底"""
    import csv
    try:
        if filename.endswith(".docx"):
            return parse_docx_in_order(io.BytesIO(file_bytes))

        elif filename.endswith(".doc"):
            # .doc 格式（旧版 Word）暂不支持，返回空字符串提示
            print(f"[{filename}] .doc 格式暂不支持，请转换为 .docx 格式")
            return ""

        elif filename.endswith(".xlsx"):
            wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
            content = []
            for sheet in wb.worksheets:
                content.append(f"\n### {sheet.title}\n")
                for i, r in enumerate(sheet.iter_rows(values_only=True)):
                    row_data = [str(c).replace('\n', ' ').replace('\r', '').strip() if c is not None else "" for c in r]
                    if not any(row_data): continue
                    content.append("| " + " | ".join(row_data) + " |")
                    if i == 0:
                        content.append("|" + "|".join(["---"] * len(row_data)) + "|")
            return "\n".join(content)

        elif filename.endswith(".csv"):
            text = file_bytes.decode('utf-8-sig', errors='ignore')
            reader = csv.reader(io.StringIO(text))
            content = []
            for i, row in enumerate(reader):
                row_data = [c.replace('\n', ' ').replace('\r', '').strip() for c in row]
                if not any(row_data): continue
                content.append("| " + " | ".join(row_data) + " |")
                if i == 0:
                    content.append("|" + "|".join(["---"] * len(row_data)) + "|")
            return "\n".join(content)

        elif filename.endswith(".pdf"):
            result_text, requires_ocr = parse_pdf_in_order(io.BytesIO(file_bytes))

            # 三重拦截机制触发云端 TextIn OCR 兜底：
            # A: 字数极少 / B: 无原生表格 / C: 包含大面积实拍图片
            if len(result_text) < 50 or "|" not in result_text or requires_ocr:
                print(f"[{filename}] 识别为混合型/扫描型 PDF，智能切换至 TextIn 云端解析...")
                return extract_text_via_textin(file_bytes)

            return result_text

    except Exception as e:
        print(f"数据解析遭遇异常: {e}")
        pass
    return ""

def loose_match_sub_element(major_element: str, file_name: str) -> str:
    """
    修改后：仅依靠上传文件名称与 JSON 进行匹配。
    修复了带日期后缀（点号）的文件名被错误截断的 Bug，
    并融合了包含、核心词根、最长公共子串、相似度四重匹配策略。
    """
    element_data = routing_dict.get(major_element, {})
    sub_elements = element_data.get("子要素", {})

    def clean_text(text: str) -> str:
        """初级清洗：剔除前缀和无用符号"""
        t = text.replace("废弃物", "废物")
        return re.sub(r'[A-Za-z0-9\-\.\:（）\(\)《》_]|From|海正药业|杭州|公司内部|有限公司', '', t).strip()

    def get_core_noun(text: str) -> str:
        """深度清洗：暴露出业务专有名词核心"""
        stopwords = ["管理", "规定", "制度", "程序", "记录", "台账", "清单", "检查", "维护", "评估", "报告", "装置", "洗眼", "处理", "调查", "表", "单", "流程", "分析", "情况", "计划"]
        t = text
        for w in stopwords:
            t = t.replace(w, "")
        return t.strip()

    # 【核心修复】：使用 os.path.splitext 安全移除扩展名，防止日期中的点号导致截断
    base_filename = os.path.splitext(file_name)[0]

    norm_filename = clean_text(base_filename)
    core_filename = get_core_noun(norm_filename)

    for rule_name, file_list in sub_elements.items():
        for expected_file in file_list:
            norm_expected = clean_text(expected_file)
            core_expected = get_core_noun(norm_expected)

            if not norm_expected:
                continue

            if norm_expected in norm_filename:
                return rule_name


            match = difflib.SequenceMatcher(None, norm_expected, norm_filename).find_longest_match(0, len(norm_expected), 0, len(norm_filename))
            if match.size >= 4:
                return rule_name

            similarity = difflib.SequenceMatcher(None, norm_expected, norm_filename).ratio()
            if similarity > 0.6:
                return rule_name

    # 匹配失败：不再返回"未知标准"，而是直接返回大要素的【主要素文件】的名称
    major_file_name = element_data.get("主要素文件", "该要素缺失主要素文件配置")
    return major_file_name

def calculate_inspection_stats(evidence_text: str):
    total_count = 0
    completed_count = 0
    cats = {}

    lines = evidence_text.strip().split("\n")
    for line in lines:
        line = line.strip()

        # 【核心修复】：剔除 Markdown 表格语法带来的最外层管道符 '|'
        # 把 "| 1 | 2026... |" 还原成干净的 "1 | 2026..."
        clean_line = line.strip("|").strip()

        if not clean_line:
            continue

        # 提取列，兼顾 Excel 解析的 ' | ' 和 CSV 的 ','
        delimiter = " | " if " | " in clean_line else ","
        cols = clean_line.split(delimiter)

        # 只要第一列（序号）不是纯数字，统统过滤掉
        # 此时 cols[0] 已经是干净的纯数字了
        if not cols or not cols[0].strip().replace('"', '').isdigit():
            continue

        total_count += 1

        # 统计是否完成
        if "已完成" in clean_line:
            completed_count += 1

        # 提取隐患类别 (通常在第4列，即索引3)
        if len(cols) > 3:
            cat = cols[3].strip().replace('"', '')
            if cat and cat not in ["", "-", "无", "N/A"]:
                cats[cat] = cats.get(cat, 0) + 1

    # 排序并取前三名
    top3 = sorted(cats.items(), key=lambda x: x[1], reverse=True)[:3]
    rate = round((completed_count / total_count * 100), 2) if total_count > 0 else 0

    return total_count, completed_count, rate, top3

def call_deepseek_daily_inspection(evidence_text: str, standard_text: str) -> dict:
    logger.info("=" * 60)
    logger.info("开始调用 DeepSeek 日常检查分析模型...")
    logger.info(f"提示词来源: 数据库 (daily_inspection)")
    logger.info(f"证据文本长度: {len(evidence_text)} 字符")

    total, completed, rate, top3 = calculate_inspection_stats(evidence_text)

    top3_details = []
    for cat, count in top3:
        percentage = round((count / total * 100), 1) if total > 0 else 0
        top3_details.append(f"{cat}（{count}项，占{percentage}%）")
    top3_str = "、".join(top3_details) if top3_details else "无明显类别"

    logger.info(f"统计信息: 总数={total}, 已完成={completed}, 整改率={rate}%")
    logger.info(f"前三类别: {top3_str}")

    base_prompt = get_prompt_from_db('daily_inspection')
    prompt = base_prompt.replace("{standard_text}", standard_text)\
                        .replace("{evidence_text}", evidence_text)\
                        .replace("{total}", str(total))\
                        .replace("{completed}", str(completed))\
                        .replace("{rate}", str(rate))\
                        .replace("{top3_str}", top3_str)

    try:
        logger.info("正在调用 deepseek-chat 模型...")
        start_time = time.time()
        response = deepseek_client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2
        )
        elapsed = time.time() - start_time
        result_text = response.choices[0].message.content.strip()
        if result_text.startswith("```json"):
            result_text = result_text[7:-3].strip()
        elif result_text.startswith("```"):
            result_text = result_text[3:-3].strip()
        logger.info(f"DeepSeek 模型调用成功! 耗时: {elapsed:.2f} 秒")
        logger.info(f"返回内容长度: {len(result_text)} 字符")
        logger.info("=" * 60)
        return json.loads(result_text)
    except Exception as e:
        logger.error(f"DeepSeek 模型调用失败: {str(e)}")
        logger.info("=" * 60)
        return {"结论": "分析出错", "逻辑缺陷描述": "大模型解析失败", "引用依据": "无", "修改意见": str(e)}

def call_deepseek_audit(standard_text: str, evidence_text: str) -> dict:
    logger.info("=" * 60)
    logger.info("开始调用 DeepSeek 审核模型...")
    logger.info(f"提示词来源: 数据库 (standard_audit)")
    logger.info(f"标准文本长度: {len(standard_text)} 字符")
    logger.info(f"证据文本长度: {len(evidence_text)} 字符")

    base_prompt = get_prompt_from_db('standard_audit')
    # 替换占位符
    prompt = base_prompt.replace("{standard_text}", standard_text).replace("{evidence_text}", evidence_text)

    try:
        logger.info("正在调用 deepseek-chat 模型...")
        start_time = time.time()
        response = deepseek_client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1 # 保持较低温度，确保逻辑推演的严谨性和稳定性
        )
        elapsed = time.time() - start_time
        result_text = response.choices[0].message.content.strip()
        if result_text.startswith("```json"):
            result_text = result_text[7:-3].strip()
        elif result_text.startswith("```"):
            result_text = result_text[3:-3].strip()
        logger.info(f"DeepSeek 模型调用成功! 耗时: {elapsed:.2f} 秒")
        logger.info(f"返回内容长度: {len(result_text)} 字符")
        logger.info("=" * 60)
        return json.loads(result_text)
    except Exception as e:
        logger.error(f"DeepSeek 模型调用失败: {str(e)}")
        logger.info("=" * 60)
        return {"结论": "审核出错", "逻辑缺陷描述": "大模型返回格式解析失败", "引用依据": "无", "修改意见": str(e)}
    
def save_report_to_db(filename, audit_mode, major_element, matched_rule, word_count, conclusion, result_json):
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO audit_reports (filename, audit_mode, major_element, matched_rule, word_count, conclusion, result_json, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    ''', (filename, audit_mode, major_element, matched_rule, word_count, conclusion, json.dumps(result_json, ensure_ascii=False), datetime.now(CHINA_TZ).strftime("%Y-%m-%d %H:%M:%S")))
    conn.commit()
    conn.close()

@app.post("/api/audit")
async def audit_files(
    major_element: str = Form(""),
    audit_mode: str = Form("standard"),
    files: List[UploadFile] = File(...)
):
    logger.info("=" * 60)
    logger.info(f"收到文件审核请求: mode={audit_mode}, element={major_element}, files={len(files)} 个")
    results = []

    for file in files:
        file_bytes = await file.read()
        unique_id = uuid.uuid4().hex[:8] # 生成8位绝对不重复的随机字母数字组合
        safe_filename = f"{unique_id}_{file.filename}"
        save_path = os.path.join(UPLOAD_DIR, safe_filename)
        with open(save_path, "wb") as buffer:
            buffer.write(file_bytes)
        file_preview_path = f"/uploads/{urllib.parse.quote(safe_filename)}"

        if file.filename.lower().endswith((".jpg", ".jpeg", ".png")):
            evidence_text = extract_text_via_textin(file_bytes) # 请确保此函数存在
        else:
            evidence_text = extract_upload_file_text(file_bytes, file.filename) # 请确保此函数存在

        if "解析失败" in evidence_text or len(evidence_text.strip()) < 5:
            results.append({
                "文件名": file.filename,
                "匹配标准": "解析失败",
                "提取字数": len(evidence_text),
                "审核结果": {
                    "结论": "不符合",
                    "逻辑缺陷描述": "文件解析失败或提取内容过短，无法判断",
                    "引用依据": "无",
                    "修改意见": "请检查文件是否损坏，或图像是否清晰。"
                },
                "原文": evidence_text,
                "原文件在线地址": file_preview_path # <--- 新增
            })
            continue

        if audit_mode == "daily_inspection":
            # 强制归于 8.1 要素，并读取 8.1 主要素文件
            target_element = "8.1要素"
            major_file_name = routing_dict.get(target_element, {}).get("主要素文件", "EHSMS-8.1运行控制")
            standard_text = get_standard_document_content(target_element, major_file_name)

            # 将标准文本传给 LLM
            llm_result = call_deepseek_daily_inspection(evidence_text, standard_text)

            results.append({
                "文件名": file.filename,
                "要素": target_element, # <-- 统一使用 8.1要素
                "审核文件": "日常检查台账分析",
                "提取字数": len(evidence_text),
                "是否符合": True,
                "审核结果": llm_result,
                "原文": evidence_text,
                "原文件在线地址": file_preview_path # <-- 新增原文留存
            })
        else:
            matched_rule = loose_match_sub_element(major_element, file.filename)
            standard_text = get_standard_document_content(major_element, matched_rule)
            llm_result = call_deepseek_audit(standard_text, evidence_text)

            is_pass = (llm_result.get("结论") == "符合")
            results.append({
                "文件名": file.filename, "要素": major_element, "审核文件": matched_rule, "提取字数": len(evidence_text),
                "是否符合": is_pass,
                "审核结果": llm_result,
                "原文": evidence_text,
                "原文件在线地址": file_preview_path # <-- 新增原文留存
            })

    logger.info(f"审核完成! 共处理 {len(results)} 个文件")
    logger.info("=" * 60)
    return {"status": "success", "data": results}

# --- 数据模型定义 ---
class PromptUpdate(BaseModel):
    content: str

class AIImproveRequest(BaseModel):
    current_prompt: str
    issues_feedback: str


# --- 提示词管理接口 ---
@app.get("/api/admin/prompts")
async def get_all_prompts():
    conn = sqlite3.connect(DB_FILE)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM system_prompts')
    rows = cursor.fetchall()
    conn.close()
    return {"status": "success", "data": [dict(r) for r in rows]}

@app.put("/api/admin/prompts/{prompt_key}")
async def update_prompt(prompt_key: str, payload: PromptUpdate):
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute("UPDATE system_prompts SET content = ?, updated_at = ? WHERE prompt_key = ?",
                   (payload.content, datetime.now(CHINA_TZ).strftime("%Y-%m-%d %H:%M:%S"), prompt_key))
    conn.commit()
    conn.close()
    return {"status": "success"}

# --- AI 提示词自动进化接口 ---
@app.post("/api/admin/prompts/improve")
async def improve_prompt_via_ai(payload: AIImproveRequest):
    logger.info("=" * 60)
    logger.info("开始调用智谱 AI 提示词优化...")
    logger.info(f"模型: glm-4.7")
    logger.info(f"用户反馈: {payload.issues_feedback[:100]}...")

    meta_prompt = f"""
    你是一个顶级的 Prompt Engineering（提示词工程）专家。
    目前系统正在使用以下提示词进行 ISO 体系文档的 AI 审核：

    【当前提示词】：
    {payload.current_prompt}

    业务人员在使用过程中，发现了以下问题或提出了如下改进需求：
    【用户反馈与痛点】：
    {payload.issues_feedback}

    请你帮我重新编写并优化这个提示词，解决用户反馈的问题。
    注意：
    1. 必须保留原有的占位符变量（如 {{standard_text}}, {{evidence_text}} 等），不要修改它们的名字。
    2. 强化系统角色设定和逻辑严密性。
    3. 只需要返回优化后的纯提示词文本，不要包含任何 "```" 代码块标记，也不要解释你的修改过程。
    """
    try:
        logger.info("正在调用智谱 AI glm-4.7 模型...")
        start_time = time.time()
        response = zhipu_client.chat.completions.create(
            model="glm-4.7",
            messages=[{"role": "user", "content": meta_prompt}],
            temperature=0.3
        )
        elapsed = time.time() - start_time
        optimized_prompt = response.choices[0].message.content.strip()
        logger.info(f"智谱 AI 模型调用成功! 耗时: {elapsed:.2f} 秒")
        logger.info(f"优化后提示词长度: {len(optimized_prompt)} 字符")
        logger.info("=" * 60)
        return {"status": "success", "data": {"optimized_prompt": optimized_prompt}}
    except Exception as e:
        logger.error(f"智谱 AI 模型调用失败: {str(e)}")
        logger.info("=" * 60)
        return {"status": "error", "message": str(e)}

# --- 字典管理接口 (简化为读写整个 JSON) ---
@app.get("/api/admin/dictionary")
async def get_dictionary():
    with open("element_routing_dictionary_final.json", "r", encoding="utf-8") as f:
        return {"status": "success", "data": json.load(f)}

@app.post("/api/admin/dictionary")
async def save_dictionary(payload: dict):
    # 动态更新全局变量和文件
    global routing_dict
    routing_dict = payload
    with open("element_routing_dictionary_final.json", "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=4)
    return {"status": "success"}

# 定义保存请求的数据模型
class BatchReportSave(BaseModel):
    batch_name: str
    total_count: int
    pass_count: int
    pass_rate: float
    details: list

# 2. 一次性保存整个队列批次的接口
@app.post("/api/reports/save")
async def save_batch_report(report: BatchReportSave):
    logger.info(f"保存报告: batch_name={report.batch_name}, total={report.total_count}份, 符合率={report.pass_rate}%")
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO batch_reports (batch_name, total_count, pass_count, pass_rate, details_json, created_at)
        VALUES (?, ?, ?, ?, ?, ?)
    ''', (report.batch_name, report.total_count, report.pass_count, report.pass_rate, json.dumps(report.details, ensure_ascii=False), datetime.now(CHINA_TZ).strftime("%Y-%m-%d %H:%M:%S")))
    conn.commit()
    conn.close()
    return {"status": "success"}

# 3. 报告中心查询接口
@app.get("/api/reports")
async def get_reports():
    logger.info("查询报告列表")
    conn = sqlite3.connect(DB_FILE)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM batch_reports ORDER BY created_at DESC')
    rows = cursor.fetchall()
    conn.close()
    report_count = len(rows)
    logger.info(f"查询报告列表完成: 共 {report_count} 份报告")
    return {"status": "success", "data": [dict(r) for r in rows]}

@app.delete("/api/reports/{report_id}")
async def delete_report(report_id: int):
    logger.info(f"删除报告: report_id={report_id}")
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute('DELETE FROM batch_reports WHERE id = ?', (report_id,))
    if cursor.rowcount == 0:
        logger.warning(f"删除报告失败: report_id={report_id} 不存在")
        conn.close()
        return {"status": "error", "message": "报告不存在"}
    conn.commit()
    conn.close()
    logger.info(f"删除报告成功: report_id={report_id}")
    return {"status": "success"}

from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

@app.get("/api/reports/export/word/{report_id}")
async def export_report_word(report_id: int):
    conn = sqlite3.connect(DB_FILE)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM batch_reports WHERE id = ?', (report_id,))
    row = cursor.fetchone()
    conn.close()

    if not row: raise HTTPException(status_code=404, detail="未找到报告")

    details = json.loads(row["details_json"])

    # 计算每个要素的统计数据
    element_stats = {}
    for item in details:
        element = item.get('要素', '未知要素')
        if element not in element_stats:
            element_stats[element] = {'total': 0, 'pass': 0, 'fail': 0}
        element_stats[element]['total'] += 1
        if item.get('是否符合', False):
            element_stats[element]['pass'] += 1
        else:
            element_stats[element]['fail'] += 1

    # 按要素总数排序
    sorted_elements = sorted(element_stats.items(), key=lambda x: x[1]['total'], reverse=True)

    doc = docx.Document()

    # --- 1. 强制全局中文字体渲染，彻底杜绝发虚和乱码 ---
    doc.styles['Normal'].font.name = u'微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    doc.styles['Normal'].font.size = Pt(11)

    # --- 2. 报告主标题 ---
    title = doc.add_heading('ISO 体系智能审核批次报告', level=0)
    title.alignment = 1 # 居中对齐

    # --- 3. 各要素审核统计（放在前面） ---
    doc.add_heading('一、 各要素审核统计', level=1)

    element_table = doc.add_table(rows=len(sorted_elements) + 1, cols=5)
    element_table.style = 'Table Grid'

    # 表头
    headers = ['要素名称', '文件总数', '符合数量', '不符合数量', '符合率']
    for i, header in enumerate(headers):
        element_table.cell(0, i).text = header
        element_table.cell(0, i).paragraphs[0].runs[0].bold = True

    # 数据行
    for row_idx, (element_name, stats) in enumerate(sorted_elements, start=1):
        rate = round((stats['pass'] / stats['total'] * 100), 2) if stats['total'] > 0 else 0
        rate_cell = element_table.cell(row_idx, 4).paragraphs[0]
        rate_run = rate_cell.add_run(f"{rate}%")

        element_table.cell(row_idx, 0).text = element_name
        element_table.cell(row_idx, 1).text = str(stats['total'])
        element_table.cell(row_idx, 2).text = str(stats['pass'])
        element_table.cell(row_idx, 3).text = str(stats['fail'])
        rate_run.bold = True
        rate_run.font.color.rgb = RGBColor(0, 128, 0) if rate >= 80 else RGBColor(204, 0, 0)

    doc.add_paragraph("\n") # 增加段落留白

    # --- 4. 批次概览信息表格 ---
    doc.add_heading('二、 批次审核概况', level=1)
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.style = 'Table Grid' # 加上标准黑色边框

    # 填充概览数据
    summary_table.cell(0,0).text = "批次名称"
    summary_table.cell(0,1).text = str(row['batch_name'])
    summary_table.cell(1,0).text = "审核时间"
    summary_table.cell(1,1).text = str(row['created_at'])
    summary_table.cell(2,0).text = "文件总数"
    summary_table.cell(2,1).text = f"{row['total_count']} 份"
    summary_table.cell(3,0).text = "综合符合率"

    # 给符合率单独上色加粗
    rate_cell = summary_table.cell(3,1).paragraphs[0]
    run = rate_cell.add_run(f"{row['pass_rate']}%  (符合 {row['pass_count']} 份)")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0) if row['pass_rate'] >= 80 else RGBColor(204, 0, 0)

    doc.add_paragraph("\n") # 增加段落留白

    # --- 5. 详细审核记录（每个文件生成一个专属表格） ---
    doc.add_heading('三、 详细审核记录', level=1)

    for i, item in enumerate(details):
        # 小节标题：文件名
        doc.add_heading(f"{i+1}. {item.get('文件名', '未知文件')}", level=2)

        # 文件的基本属性
        p = doc.add_paragraph()
        p.add_run(f"【所属要素】: ").bold = True
        p.add_run(f"{item.get('要素', '未知')}    ")
        p.add_run(f"【判定依据】: ").bold = True
        p.add_run(f"{item.get('审核文件', '未知')}")

        res = item.get('审核结果', {})

        # 绘制该文件的详细剖析表格
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'

        # 固定左侧列宽，让排版更美观
        table.columns[0].width = Cm(3.5)
        table.columns[1].width = Cm(12.5)

        # 第一行：结论
        table.cell(0, 0).text = "1. 审核结论"
        table.cell(0, 1).text = str(res.get("结论", "无"))
        table.cell(0, 1).paragraphs[0].runs[0].bold = True # 结论加粗

        # 第二行：问题剖析
        table.cell(1, 0).text = "2. 核心问题剖析"
        table.cell(1, 1).text = str(res.get("逻辑缺陷描述", "无"))

        # 第三行：引用依据
        table.cell(2, 0).text = "3. 引用依据"
        table.cell(2, 1).text = str(res.get("引用依据", "无"))

        # 第四行：修改意见
        table.cell(3, 0).text = "4. 修改意见"
        table.cell(3, 1).text = str(res.get("修改意见", "无"))
        # 给修改意见换个颜色
        if table.cell(3, 1).paragraphs[0].runs:
            table.cell(3, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 80, 150)

        doc.add_paragraph("\n") # 每个文件后增加留白

    # --- 4. 内存流保存与安全编码下发 ---
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    # 绝对安全的 UTF-8 文件名编码
    filename = f"ISO审核批次报告_{row['id']}.docx"
    encoded_filename = quote(filename)
    headers = {
        'Content-Disposition': f"attachment; filename*=UTF-8''{encoded_filename}"
    }
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )


@app.get("/api/reports/export/pdf/{report_id}")
async def export_report_pdf(report_id: int):
    """
    后端生成 PDF 报告（使用 reportlab 纯 Python）
    不依赖 Word，直接生成 PDF，保证内容完整不遗漏
    """
    import io
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from urllib.parse import quote

    conn = sqlite3.connect(DB_FILE)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM batch_reports WHERE id = ?', (report_id,))
    row = cursor.fetchone()
    conn.close()

    if not row:
        raise HTTPException(status_code=404, detail="未找到报告")

    details = json.loads(row["details_json"])

    # 计算每个要素的统计数据
    element_stats = {}
    for item in details:
        element = item.get('要素', '未知要素')
        if element not in element_stats:
            element_stats[element] = {'total': 0, 'pass': 0, 'fail': 0}
        element_stats[element]['total'] += 1
        if item.get('是否符合', False):
            element_stats[element]['pass'] += 1
        else:
            element_stats[element]['fail'] += 1

    sorted_elements = sorted(element_stats.items(), key=lambda x: x[1]['total'], reverse=True)

    # 注册中文字体 - 使用系统字体
    font_path = "C:/Windows/Fonts/msyh.ttc"  # 微软雅黑
    bold_font_path = "C:/Windows/Fonts/msyhbd.ttc"  # 微软雅黑 bold
    if not os.path.exists(font_path):
        font_path = "C:/Windows/Fonts/simsun.ttc"  # 宋体
        bold_font_path = "C:/Windows/Fonts/simhei.ttf"  # 黑体

    # 注册普通字体和粗体字体
    pdfmetrics.registerFont(TTFont("SimSun", font_path))
    pdfmetrics.registerFont(TTFont("SimSun-Bold", bold_font_path))

    # 创建内存 PDF
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2*cm,
        leftMargin=2*cm,
        topMargin=2*cm,
        bottomMargin=2*cm
    )

    # 样式定义
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontName="SimSun-Bold",
        fontSize=16,
        textColor=colors.Color(0.117, 0.25, 0.686),
        spaceAfter=10,
        alignment=1  # 居中
    )
    heading2_style = ParagraphStyle(
        'CustomHeading2',
        parent=styles['Heading2'],
        fontName="SimSun-Bold",
        fontSize=12,
        textColor=colors.Color(0.23, 0.51, 0.96),
        spaceAfter=8,
        spaceBefore=12
    )
    normal_style = ParagraphStyle(
        'Normal',
        parent=styles['Normal'],
        fontName="SimSun",
        fontSize=10,
        leading=14
    )

    # 流程元素列表
    elements = []

    # 标题
    elements.append(Paragraph("ISO 体系智能审核批次报告", title_style))
    elements.append(Paragraph(" ", normal_style))

    # 各要素审核统计
    elements.append(Paragraph("一、各要素审核统计", heading2_style))

    # 表头
    element_data = [["要素名称", "文件总数", "符合数量", "不符合数量", "符合率"]]
    for element_name, stats in sorted_elements:
        rate = round((stats['pass'] / stats['total'] * 100), 2) if stats['total'] > 0 else 0
        rate_str = f"{rate}%"
        rate_color = "green" if rate >= 80 else "red"
        element_data.append([
            element_name,
            str(stats['total']),
            str(stats['pass']),
            str(stats['fail']),
            rate_str
        ])

    # 创建表格
    element_table = Table(element_data, colWidths=[4*cm, 3*cm, 3*cm, 3*cm, 2.5*cm])
    element_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.Color(0.953, 0.953, 0.953)),
        ('FONTNAME', (0, 0), (-1, 0), 'SimSun-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.Color(0.867, 0.867, 0.867)),
        ('FONTNAME', (0, 1), (-1, -1), 'SimSun'),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
    ]))
    elements.append(element_table)
    elements.append(Paragraph(" ", normal_style))

    # 批次审核概况
    elements.append(Paragraph("二、批次审核概况", heading2_style))

    total = row['total_count']
    pass_count = row['pass_count']
    fail_count = total - pass_count
    rate = round(row['pass_rate'], 2) if row['pass_rate'] else 0
    rate_color = "green" if rate >= 80 else "red"

    summary_data = [
        ["批次名称", row['batch_name']],
        ["审核时间", str(row['created_at'])],
        ["文件总数", f"{total} 份"],
        ["符合数量", str(pass_count)],
        ["不符合数量", str(fail_count)],
        ["综合符合率", f"<b><font color='{rate_color}'>{rate}%</font></b>"],
    ]

    summary_table = Table(summary_data, colWidths=[4*cm, 10*cm])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.Color(0.953, 0.953, 0.953)),
        ('FONTNAME', (0, 0), (-1, -1), 'SimSun'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.Color(0.867, 0.867, 0.867)),
        ('FONTNAME', (0, 0), (0, -1), 'SimSun-Bold'),
    ]))
    elements.append(summary_table)
    elements.append(Paragraph(" ", normal_style))

    # 详细审核记录
    elements.append(Paragraph("三、详细审核记录", heading2_style))

    for i, item in enumerate(details):
        filename = item.get('文件名', '未知文件')
        element = item.get('要素', '未知')
        rule = item.get('审核文件', '未知')
        is_pass = item.get('是否符合', False)
        res = item.get('审核结果', {})
        res_text = "符合" if is_pass else "不符合"
        res_color = "green" if is_pass else "red"

        # 文件名标题
        elements.append(Paragraph(f"{i+1}. {filename}", heading2_style))

        # 基本信息
        info_text = f'【所属要素】: {element}    【判定依据】: {rule}'
        elements.append(Paragraph(info_text, normal_style))

        # 结论
        elements.append(Paragraph(f'【审核结论】: <font color="{res_color}">{res_text}</font>', normal_style))

        # 问题剖析
        elements.append(Paragraph('【核心问题剖析】:', normal_style))
        description = res.get('逻辑缺陷描述', '无')
        if description:
            # 分段处理长文本
            for j in range(0, len(description), 100):
                elements.append(Paragraph(description[j:j+100], normal_style))

        # 引用依据
        elements.append(Paragraph('【引用依据】:', normal_style))
        elements.append(Paragraph(res.get('引用依据', '无'), normal_style))

        # 修改意见
        elements.append(Paragraph('【修改意见】:', normal_style))
        elements.append(Paragraph(res.get('修改意见', '无'), normal_style))
        elements.append(Paragraph(" ", normal_style))

    # 生成 PDF
    doc.build(elements)

    # 获取 PDF 内容
    buffer.seek(0)
    pdf_content = buffer.read()
    buffer.close()

    filename = f"ISO审核批次报告_{row['id']}.pdf"
    encoded_filename = quote(filename)
    headers = {
        'Content-Disposition': f"attachment; filename*=UTF-8''{encoded_filename}"
    }

    return StreamingResponse(
        io.BytesIO(pdf_content),
        media_type="application/pdf",
        headers=headers
    )


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host=BACKEND_HOST, port=BACKEND_PORT)
