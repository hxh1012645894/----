"""
报告API路由（含导出）
"""
import os
import io
import json
import sqlite3
import docx
from datetime import datetime
from urllib.parse import quote
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

from config import DB_FILE, logger, CHINA_TZ
from models import BatchReportSave

router = APIRouter(prefix="/api/reports", tags=["报告"])


@router.post("/save")
async def save_batch_report(report: BatchReportSave):
    """保存批次报告"""
    logger.info(f"保存报告: batch_name={report.batch_name}, total={report.total_count}份, 符合率={report.pass_rate}%")
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO batch_reports (batch_name, total_count, pass_count, pass_rate, details_json, created_at)
        VALUES (?, ?, ?, ?, ?, ?)
    ''', (report.batch_name, report.total_count, report.pass_count, report.pass_rate,
          json.dumps(report.details, ensure_ascii=False), datetime.now(CHINA_TZ).strftime("%Y-%m-%d %H:%M:%S")))
    conn.commit()
    conn.close()
    return {"status": "success"}


@router.get("")
async def get_reports():
    """获取报告列表"""
    logger.info("查询报告列表")
    conn = sqlite3.connect(DB_FILE)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM batch_reports ORDER BY created_at DESC')
    rows = cursor.fetchall()
    conn.close()
    report_count = len(rows)
    logger.info(f"查询报告列表完成: 共 {report_count} 份报告")
    return {"status": "success", "data": [dict(r) for r in rows]}


@router.delete("/{report_id}")
async def delete_report(report_id: int):
    """删除报告"""
    logger.info(f"删除报告: report_id={report_id}")
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute('DELETE FROM batch_reports WHERE id = ?', (report_id,))
    if cursor.rowcount == 0:
        logger.warning(f"删除报告失败: report_id={report_id} 不存在")
        conn.close()
        return {"status": "error", "message": "报告不存在"}
    conn.commit()
    conn.close()
    logger.info(f"删除报告成功: report_id={report_id}")
    return {"status": "success"}


@router.get("/export/word/{report_id}")
async def export_report_word(report_id: int):
    """导出Word报告"""
    conn = sqlite3.connect(DB_FILE)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM batch_reports WHERE id = ?', (report_id,))
    row = cursor.fetchone()
    conn.close()

    if not row:
        raise HTTPException(status_code=404, detail="未找到报告")

    details = json.loads(row["details_json"])

    # 计算要素统计
    element_stats = {}
    for item in details:
        element = item.get('要素', '未知要素')
        if element not in element_stats:
            element_stats[element] = {'total': 0, 'pass': 0, 'fail': 0}
        element_stats[element]['total'] += 1
        if item.get('是否符合', False):
            element_stats[element]['pass'] += 1
        else:
            element_stats[element]['fail'] += 1

    sorted_elements = sorted(element_stats.items(), key=lambda x: x[1]['total'], reverse=True)

    doc = docx.Document()

    # 全局中文字体
    doc.styles['Normal'].font.name = u'微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    doc.styles['Normal'].font.size = Pt(11)

    # 报告标题
    title = doc.add_heading('ISO 体系智能审核批次报告', level=0)
    title.alignment = 1

    # 各要素审核统计
    doc.add_heading('一、 各要素审核统计', level=1)
    element_table = doc.add_table(rows=len(sorted_elements) + 1, cols=5)
    element_table.style = 'Table Grid'

    headers = ['要素名称', '文件总数', '符合数量', '不符合数量', '符合率']
    for i, header in enumerate(headers):
        element_table.cell(0, i).text = header
        element_table.cell(0, i).paragraphs[0].runs[0].bold = True

    for row_idx, (element_name, stats) in enumerate(sorted_elements, start=1):
        rate = round((stats['pass'] / stats['total'] * 100), 2) if stats['total'] > 0 else 0
        rate_cell = element_table.cell(row_idx, 4).paragraphs[0]
        rate_run = rate_cell.add_run(f"{rate}%")
        element_table.cell(row_idx, 0).text = element_name
        element_table.cell(row_idx, 1).text = str(stats['total'])
        element_table.cell(row_idx, 2).text = str(stats['pass'])
        element_table.cell(row_idx, 3).text = str(stats['fail'])
        rate_run.bold = True
        rate_run.font.color.rgb = RGBColor(0, 128, 0) if rate >= 80 else RGBColor(204, 0, 0)

    doc.add_paragraph("\n")

    # 批次审核概况
    doc.add_heading('二、 批次审核概况', level=1)
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.style = 'Table Grid'

    summary_table.cell(0, 0).text = "批次名称"
    summary_table.cell(0, 1).text = str(row['batch_name'])
    summary_table.cell(1, 0).text = "审核时间"
    summary_table.cell(1, 1).text = str(row['created_at'])
    summary_table.cell(2, 0).text = "文件总数"
    summary_table.cell(2, 1).text = f"{row['total_count']} 份"
    summary_table.cell(3, 0).text = "综合符合率"

    rate_cell = summary_table.cell(3, 1).paragraphs[0]
    run = rate_cell.add_run(f"{row['pass_rate']}%  (符合 {row['pass_count']} 份)")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0) if row['pass_rate'] >= 80 else RGBColor(204, 0, 0)

    doc.add_paragraph("\n")

    # 详细审核记录
    doc.add_heading('三、 详细审核记录', level=1)

    for i, item in enumerate(details):
        doc.add_heading(f"{i+1}. {item.get('文件名', '未知文件')}", level=2)

        p = doc.add_paragraph()
        p.add_run(f"【所属要素】: ").bold = True
        p.add_run(f"{item.get('要素', '未知')}    ")
        p.add_run(f"【判定依据】: ").bold = True
        p.add_run(f"{item.get('审核文件', '未知')}")

        res = item.get('审核结果', {})

        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        table.columns[0].width = Cm(3.5)
        table.columns[1].width = Cm(12.5)

        table.cell(0, 0).text = "1. 审核结论"
        table.cell(0, 1).text = str(res.get("结论", "无"))
        table.cell(0, 1).paragraphs[0].runs[0].bold = True

        table.cell(1, 0).text = "2. 核心问题剖析"
        table.cell(1, 1).text = str(res.get("逻辑缺陷描述", "无"))

        table.cell(2, 0).text = "3. 引用依据"
        table.cell(2, 1).text = str(res.get("引用依据", "无"))

        table.cell(3, 0).text = "4. 修改意见"
        table.cell(3, 1).text = str(res.get("修改意见", "无"))
        if table.cell(3, 1).paragraphs[0].runs:
            table.cell(3, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 80, 150)

        doc.add_paragraph("\n")

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    filename = f"ISO审核批次报告_{row['id']}.docx"
    encoded_filename = quote(filename)
    headers = {'Content-Disposition': f"attachment; filename*=UTF-8''{encoded_filename}"}
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )


@router.get("/export/pdf/{report_id}")
async def export_report_pdf(report_id: int):
    """导出PDF报告"""
    conn = sqlite3.connect(DB_FILE)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM batch_reports WHERE id = ?', (report_id,))
    row = cursor.fetchone()
    conn.close()

    if not row:
        raise HTTPException(status_code=404, detail="未找到报告")

    details = json.loads(row["details_json"])

    # 计算要素统计
    element_stats = {}
    for item in details:
        element = item.get('要素', '未知要素')
        if element not in element_stats:
            element_stats[element] = {'total': 0, 'pass': 0, 'fail': 0}
        element_stats[element]['total'] += 1
        if item.get('是否符合', False):
            element_stats[element]['pass'] += 1
        else:
            element_stats[element]['fail'] += 1

    sorted_elements = sorted(element_stats.items(), key=lambda x: x[1]['total'], reverse=True)

    # 注册中文字体
    font_path = "C:/Windows/Fonts/msyh.ttc"
    bold_font_path = "C:/Windows/Fonts/msyhbd.ttc"
    if not os.path.exists(font_path):
        font_path = "C:/Windows/Fonts/simsun.ttc"
        bold_font_path = "C:/Windows/Fonts/simhei.ttf"

    pdfmetrics.registerFont(TTFont("SimSun", font_path))
    pdfmetrics.registerFont(TTFont("SimSun-Bold", bold_font_path))

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=2*cm, leftMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontName="SimSun-Bold",
                                  fontSize=16, textColor=colors.Color(0.117, 0.25, 0.686), spaceAfter=10, alignment=1)
    heading2_style = ParagraphStyle('CustomHeading2', parent=styles['Heading2'], fontName="SimSun-Bold",
                                      fontSize=12, textColor=colors.Color(0.23, 0.51, 0.96), spaceAfter=8, spaceBefore=12)
    normal_style = ParagraphStyle('Normal', parent=styles['Normal'], fontName="SimSun", fontSize=10, leading=14)

    elements = []
    elements.append(Paragraph("ISO 体系智能审核批次报告", title_style))
    elements.append(Paragraph(" ", normal_style))

    elements.append(Paragraph("一、各要素审核统计", heading2_style))
    element_data = [["要素名称", "文件总数", "符合数量", "不符合数量", "符合率"]]
    for element_name, stats in sorted_elements:
        rate = round((stats['pass'] / stats['total'] * 100), 2) if stats['total'] > 0 else 0
        element_data.append([element_name, str(stats['total']), str(stats['pass']), str(stats['fail']), f"{rate}%"])

    element_table = Table(element_data, colWidths=[4*cm, 3*cm, 3*cm, 3*cm, 2.5*cm])
    element_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.Color(0.953, 0.953, 0.953)),
        ('FONTNAME', (0, 0), (-1, 0), 'SimSun-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.Color(0.867, 0.867, 0.867)),
        ('FONTNAME', (0, 1), (-1, -1), 'SimSun'),
    ]))
    elements.append(element_table)
    elements.append(Paragraph(" ", normal_style))

    elements.append(Paragraph("二、批次审核概况", heading2_style))
    rate = round(row['pass_rate'], 2) if row['pass_rate'] else 0
    rate_color = "green" if rate >= 80 else "red"
    summary_data = [
        ["批次名称", row['batch_name']],
        ["审核时间", str(row['created_at'])],
        ["文件总数", f"{row['total_count']} 份"],
        ["符合数量", str(row['pass_count'])],
        ["不符合数量", str(row['total_count'] - row['pass_count'])],
        ["综合符合率", f"<b><font color='{rate_color}'>{rate}%</font></b>"],
    ]
    summary_table = Table(summary_data, colWidths=[4*cm, 10*cm])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.Color(0.953, 0.953, 0.953)),
        ('FONTNAME', (0, 0), (-1, -1), 'SimSun'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.Color(0.867, 0.867, 0.867)),
        ('FONTNAME', (0, 0), (0, -1), 'SimSun-Bold'),
    ]))
    elements.append(summary_table)
    elements.append(Paragraph(" ", normal_style))

    elements.append(Paragraph("三、详细审核记录", heading2_style))
    for i, item in enumerate(details):
        filename = item.get('文件名', '未知文件')
        res = item.get('审核结果', {})
        is_pass = item.get('是否符合', False)
        res_color = "green" if is_pass else "red"

        elements.append(Paragraph(f"{i+1}. {filename}", heading2_style))
        elements.append(Paragraph(f'【所属要素】: {item.get("要素", "未知")}    【判定依据】: {item.get("审核文件", "未知")}', normal_style))
        elements.append(Paragraph(f'【审核结论】: <font color="{res_color}">{"符合" if is_pass else "不符合"}</font>', normal_style))
        elements.append(Paragraph('【核心问题剖析】:', normal_style))
        for j in range(0, len(res.get('逻辑缺陷描述', '无')), 100):
            elements.append(Paragraph(res.get('逻辑缺陷描述', '无')[j:j+100], normal_style))
        elements.append(Paragraph('【引用依据】:', normal_style))
        elements.append(Paragraph(res.get('引用依据', '无'), normal_style))
        elements.append(Paragraph('【修改意见】:', normal_style))
        elements.append(Paragraph(res.get('修改意见', '无'), normal_style))
        elements.append(Paragraph(" ", normal_style))

    doc.build(elements)
    buffer.seek(0)
    pdf_content = buffer.read()
    buffer.close()

    filename = f"ISO审核批次报告_{row['id']}.pdf"
    encoded_filename = quote(filename)
    headers = {'Content-Disposition': f"attachment; filename*=UTF-8''{encoded_filename}"}
    return StreamingResponse(io.BytesIO(pdf_content), media_type="application/pdf", headers=headers)