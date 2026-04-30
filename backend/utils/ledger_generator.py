"""
台账生成工具：Word/Excel台账生成
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side

from config import UPLOAD_DIR, CHINA_TZ


def generate_accident_ledger_word(accident_data: dict, analysis_result: dict, measures: list = None) -> str:
    """
    生成Word格式的事故台账

    台账内容：
    - 事故基本信息表
    - AI分析结果
    - 整改措施列表
    """
    doc = Document()

    # 标题
    title = doc.add_heading('事故台账', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 台账编号和时间
    doc.add_paragraph(f"台账编号：{accident_data.get('id', '')}")
    doc.add_paragraph(f"生成时间：{datetime.now(CHINA_TZ).strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()

    # 一、事故基本信息
    doc.add_heading('一、事故基本信息', level=1)

    # 基本信息表格
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['项目', '内容']
    table.rows[0].cells[0].text = headers[0]
    table.rows[0].cells[1].text = headers[1]

    info_rows = [
        ('事故时间', accident_data.get('accident_time', '')),
        ('事故地点', accident_data.get('location', '')),
        ('事故类型', accident_data.get('accident_type', '')),
        ('伤亡情况', accident_data.get('casualties', '') or '无伤亡'),
        ('所属部门', accident_data.get('department', '') or '未填写'),
        ('属地工程师', accident_data.get('engineer_name', '') or '未填写'),
        ('事故描述', accident_data.get('description', '') or '无')
    ]

    for i, (label, value) in enumerate(info_rows, 1):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # 二、AI分析结果
    doc.add_heading('二、AI分析结果', level=1)

    # 直接原因
    p = doc.add_paragraph()
    p.add_run('1. 直接原因分析').bold = True
    doc.add_paragraph(analysis_result.get('direct_cause', '未分析'))

    # 间接原因
    p = doc.add_paragraph()
    p.add_run('2. 间接原因分析').bold = True
    doc.add_paragraph(analysis_result.get('indirect_cause', '未分析'))

    # 事故教训
    p = doc.add_paragraph()
    p.add_run('3. 事故教训总结').bold = True
    doc.add_paragraph(analysis_result.get('lessons_learned', '无'))

    # 整改措施建议
    p = doc.add_paragraph()
    p.add_run('4. 整改措施建议').bold = True
    doc.add_paragraph(analysis_result.get('rectification_measures', '无'))

    doc.add_paragraph()

    # 三、整改措施执行情况（如有）
    if measures and len(measures) > 0:
        doc.add_heading('三、整改措施执行情况', level=1)

        # 整改措施表格
        measure_table = doc.add_table(rows=len(measures) + 1, cols=5)
        measure_table.style = 'Table Grid'

        headers = ['序号', '整改措施', '责任人', '整改期限', '状态']
        for i, h in enumerate(headers):
            measure_table.rows[0].cells[i].text = h

        for i, m in enumerate(measures, 1):
            measure_table.rows[i].cells[0].text = str(i)
            measure_table.rows[i].cells[1].text = m.get('measure_content', '')
            measure_table.rows[i].cells[2].text = m.get('responsible_person', '') or '未指派'
            measure_table.rows[i].cells[3].text = m.get('deadline', '') or '未设置'
            measure_table.rows[i].cells[4].text = m.get('status', 'pending')

        doc.add_paragraph()

    # 保存文件
    filename = f"ledger_{accident_data.get('id', '')}_{datetime.now(CHINA_TZ).strftime('%Y%m%d%H%M%S')}.docx"
    filepath = os.path.join(UPLOAD_DIR, filename)
    doc.save(filepath)

    return f"/uploads/{filename}"


def generate_batch_ledger_excel(accidents: list) -> str:
    """
    批量生成Excel格式的台账汇总表
    """
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = '事故台账汇总'

    # 样式定义
    header_font = Font(bold=True, size=12)
    header_alignment = Alignment(horizontal='center', vertical='center')
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    # 表头
    headers = ['序号', '事故时间', '事故类型', '事故地点', '伤亡情况', '所属部门',
               '属地工程师', '直接原因', '间接原因', '整改措施', '状态']

    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.alignment = header_alignment
        cell.border = thin_border

    # 设置列宽
    column_widths = [6, 18, 12, 20, 10, 15, 12, 40, 40, 40, 10]
    for i, width in enumerate(column_widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = width

    # 数据行
    for row, accident in enumerate(accidents, 2):
        ws.cell(row=row, column=1, value=row - 1).border = thin_border
        ws.cell(row=row, column=2, value=accident.get('accident_time', '')).border = thin_border
        ws.cell(row=row, column=3, value=accident.get('accident_type', '')).border = thin_border
        ws.cell(row=row, column=4, value=accident.get('location', '')).border = thin_border
        ws.cell(row=row, column=5, value=accident.get('casualties', '') or '无').border = thin_border
        ws.cell(row=row, column=6, value=accident.get('department', '') or '').border = thin_border
        ws.cell(row=row, column=7, value=accident.get('engineer_name', '') or '').border = thin_border
        ws.cell(row=row, column=8, value=accident.get('direct_cause', '') or '').border = thin_border
        ws.cell(row=row, column=9, value=accident.get('indirect_cause', '') or '').border = thin_border
        ws.cell(row=row, column=10, value=accident.get('rectification_measures', '') or '').border = thin_border

        status_cell = ws.cell(row=row, column=11, value=accident.get('status', 'draft'))
        status_cell.border = thin_border

        # 状态颜色标记
        if accident.get('status') == 'submitted':
            status_cell.font = Font(color='006400')  # 绿色
        else:
            status_cell.font = Font(color='808080')  # 灰色

    # 保存文件
    filename = f"ledger_batch_{datetime.now(CHINA_TZ).strftime('%Y%m%d%H%M%S')}.xlsx"
    filepath = os.path.join(UPLOAD_DIR, filename)
    wb.save(filepath)

    return f"/uploads/{filename}"